import logging
import os
import signal
import sys
from datetime import datetime

import docx.enum.text
import docx2txt
import nltk
import spacy
from docx import Document
from pyfiglet import Figlet



# updatedUnncertain Tag infixde
# pass in updatedUnccertain tags [auto tagging]
# lloook
def testFileWrite():
    str = "hey alyssa today's a pretty beautiful day."
    arr = str.split(" ")
    destinationFileName = "./Outputs/highlight_testing.docx"
    document = Document()
    p = document.add_paragraph("")
    for idx in range(len(arr)):
        word = arr[idx]
        if(word == "alyssa") :
            # run.font.highlight_color = docx.enum.text.WD_COLOR.YELLOW
            p.add_run(word).font.highlight_color = docx.enum.text.WD_COLOR.YELLOW
        else:
            p.add_run(word)
        p.add_run(" ")

    # we need to write the contents of the arr into a docx, just write out first, then we look at how to highlight:
    # document.add_heading('Document Title', 0)
    document.save(destinationFileName)


"""

    sourceFileName = "./Inputs/debug.docx"
    target = "fuck"
    # test out the instructions here first
    doc = Document(sourceFileName)
    for paragraph in doc.paragraphs:
        print(f"\n this is what the paragraph is: \n{paragraph}")
        for run in paragraph.runs:
            print(f"\n this is what run.text gives: \n{run.text}")
        if target in paragraph.text:
            for run in paragraph.runs:
                print(f"\n this is what run.text gives: \n{run.text}")
                #if target in run.text:
                   #  x = run.text.split(target)
                #    run.clear()
                    #for i in range(len(x) - 1):
                    #    run.add_text(x[i])
                    #    run.add_text(target)
                    #    run.font.highlight_color = docx.enum.text.WD_COLOR.YELLOW

    # doc.save("fuckme.docx")

"""


"""
def Alyssa_foo():
    sourceFileName = "./Inputs/debug.docx"
    # target = "fuck"
    # test out the instructions here first
    doc = Document(sourceFileName)
    for paragraph in doc.paragraphs:
        print(f"\n this is what the paragraph is: \n{paragraph}")
        for run in paragraph.runs:
            print(f"\n this is what run.text gives: \n{run.text}")
        if target in paragraph.text:
            for run in paragraph.runs:
                print(f"\n this is what run.text gives: \n{run.text}")
                # if target in run.text:
                #  x = run.text.split(target)
                #    run.clear()
                # for i in range(len(x) - 1):
                #    run.add_text(x[i])
                #    run.add_text(target)
                #    run.font.highlight_color = docx.enum.text.WD_COLOR.YELLOW

    # doc.save("fuckme.docx")
    return

"""


def main():
    print("hello world")
    testFileWrite()



if __name__ == "__main__":
    main()

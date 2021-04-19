""" ==========================================
This is a convenience program written to improve our 
manual checking of PoS tagging done. This should serve
as a useful boilerplate to add in more functionality
as a user so wishes. 

Next versions will cross check existing POS tagging libs
and ask the human only if there are discrepancies.

possible issues to face: 
1. the workspace file is overwritten to reflect that most 
   status. i.e. the unchecked terms. but the overwriting
   does not preserve the original whitespace in the source
   document. (actually we never ever preserve the whitespace
   from the source document)

Questions:
1. Input file related queries:
    a. How come there are duplicate words in the input file?
    b. How come some words are erroneously joined "generalhospital"?

Todos:
0. pass in all the different text files <--- handle multiple files
1. Fix the duplicated words problem
   - see alyssa's use of the enchant dictionary to filter through for words not in a dictionary

Written by: Alyssa Nah Xiao Ting and Ritesh Kumar
 ============================================== """
import logging
import os
import signal
import sys
from datetime import datetime

import docx.enum.text
import docx2txt
import nltk
import spacy
from docx import Document
from pyfiglet import Figlet

logger = logging.getLogger("logger")
logging.basicConfig(filename="logging_output.txt",
                    filemode='a+',
                    format='%(asctime)s,%(msecs)d %(name)s %(levelname)s %(message)s',
                    datefmt='%H:%M:%S',
                    level=logging.DEBUG)

# define constants here
global THRESHOLD
NUM_MODELS_USED = 2 # change this based on how many other models are being used.
NUMBER_OF_UNCERTAINTIES = None
f = Figlet(font="colossal")
# TODO: import this wordlist from somewhere in the internet. there should be singlish wordlists
SinglishWords = ["LA", "LAH", "LOR", "AH", "MEH", "LEH", "HOR"]
PseudoSinglishWords = ["ONE", "WHAT"]
tagsChanged = []
SINGLISH_SCORE = -1
PSEUDOSINGLISH_SCORE = 0
SINGLISH_HIGHLIGHT_COLOUR = docx.enum.text.WD_COLOR.TURQUOISE
PSEUDOSINGLISH_HIGHLIGHT_COLOUR = docx.enum.text.WD_COLOR.PINK
UNCERTAIN_HIGHLIGHT_COLOUR = docx.enum.text.WD_COLOR.YELLOW
DESTINATION_FOLDER = os.path.join(os.getcwd(), "Outputs")
global DESTINATION_FILE_PATH
INPUT_FOLDER = os.path.join(os.getcwd(), "Inputs")

def main():
    logger.debug("\n\n\n>>>>>>>>>>>>>>>>>>>>>> Running Script Now <<<<<<<<<<<<<<<<<<<<<<<<<<<<<\n\n\n")
    desiredPercentage = float(sys.argv[1])
    global THRESHOLD
    global DESTINATION_FILE_PATH
    global DESTINATION_FOLDER
    THRESHOLD = (desiredPercentage / 100.0) * NUM_MODELS_USED
    filePaths = getInputFilepaths();

    for filePath in filePaths:
        sourceFileName = os.path.basename(filePath)
        fileTitle = sourceFileName.split(".")[0]
        workspaceFileName = fileTitle + "_workspace.txt"
        outputFileName = fileTitle + ".docx"
        DESTINATION_FILE_PATH = os.path.join(DESTINATION_FOLDER, outputFileName)
        contents = init(sourceFileName, workspaceFileName)
        try:
            checkPOS(contents)
        except Exception as e:
            print(f" This file is fishy: {fileTitle} because: \n {e}")

    '''
        # note that the input file has to be in the same project directory
       sourceFileName = sys.argv[1]
       desiredPercentage = float(sys.argv[2])
       global THRESHOLD
       THRESHOLD = (desiredPercentage / 100.0) * NUM_MODELS_USED
       fileTitle = sourceFileName.split(".")[0]
       workspaceFileName = fileTitle + "_workspace.txt"
       contents = init(sourceFileName, workspaceFileName)
       #finalWords, finalTags, stats = checkPOS(contents)
       checkPOS(contents)

       '''



    """
    writeToOutputFile(fileTitle, finalWords, finalTags)
    print(endingGreeting + stats)
    logger.debug("\n\n\n>>>>>>>>>>>>>>>>>>>>>> Done Running script <<<<<<<<<<<<<<<<<<<<<<<<<<<<<\n\n\n")
    # note that the input file has to be in the same project directory
    """

# returns an array containing tagging done by models available by spacy
def generateSpacyTags(words):
    prose = ""
    for word in words:
        prose += word + " "
    prose = prose[:-1] # to remove the extra space at the end of the prose.
    taggerSM, taggerMD, taggerLG = spacy.load("en_core_web_sm"), spacy.load("en_core_web_md"), spacy.load("en_core_web_lg")
    sm_doc, md_doc, lg_doc = taggerSM(prose), taggerMD(prose), taggerLG(prose)
    sm_tags, md_tags, lg_tags = [], [], []

    # concat the tags, only need to handle the tags array.
    sm_idx, word_idx = 0 , 0
    generatedTokens = []
    # while(word_idx < len(words) and sm_idx < len(sm_doc)) :
    while (sm_idx < len(sm_doc)):
        currentToken = sm_doc[sm_idx]
        isFirstOrLast = sm_idx == len(sm_doc) or 0
        if(currentToken.text == "-" and not isFirstOrLast): #  .... last few words -
            prevToken, nextToken = sm_doc[sm_idx - 1] , sm_doc[sm_idx + 1]
            concatenatedWord = prevToken.text + currentToken.text + nextToken.text
            nltkTag = nltk.pos_tag(nltk.word_tokenize(concatenatedWord))[0][1]
            sm_tags.pop(); md_tags.pop(); lg_tags.pop()
            generatedTokens.pop()
            generatedTokens.append(concatenatedWord)
            sm_tags.append(nltkTag), md_tags.append(nltkTag), lg_tags.append(nltkTag)
            sm_idx += 1
        else:
            generatedTokens.append(sm_doc[sm_idx].text)
            sm_tags.append(sm_doc[sm_idx].tag_)
            md_tags.append(md_doc[sm_idx].tag_)
            lg_tags.append(lg_doc[sm_idx].tag_)
        sm_idx += 1; word_idx += 1  # increment both pointers by 1 for the next iter of while loop
    assert (len(words) == len(sm_tags) == len(md_tags) == len(lg_tags))
    return (sm_tags, md_tags, lg_tags)

# returns an array containing NLTK tags
def generateNLTKtags(words):
    nltkOutput = nltk.pos_tag(words)
    nltkTags = []
    for term in nltkOutput:
        nltkTags.append(term[1])  # nb: nltk output is an array of tuples that's why nltkOutput: [(<word>_<tag>), etc..]
    return nltkTags

def getSourceContent(sourceUrl):
    if(".docx" in sourceUrl):
        return docx2txt.process(sourceUrl)
    elif (".txt" in sourceUrl):
        # TODO: the contractions have an extra backslash, check if this affects anything or can it be ignored.
        #       because in thefirst place we are ignoring contractions. This is suspected to be the reason why
        #       our assertions in other areas fail (this is the suspected bug) and can be avoided if we get
        #       the original files that have not been preprocessed by stanford.
        fo = open(sourceUrl, "r")
        content = fo.read().replace("\n", " ").replace("\'", " ")
        return content
    else:
         assert False, "no idea what file type this is"

# returns contents from the workspace if workspace aldy exists, else copies from source to new workspace. This is
# to prevent modifying the source directly:
def init(sourceFileName, workspaceFileName):
    sourceUrl = os.path.join(os.getcwd(),"Inputs", sourceFileName)
    workspaceUrl = os.path.join(os.getcwd(),"Workspaces", workspaceFileName)
    if os.path.isfile(workspaceUrl):
        # logger.debug("workspace file exists, shall read from it")
        workspaceContent = open(workspaceUrl, "r").read()
    else:
        sourceContent = getSourceContent(sourceUrl)
        workspaceContent = writeToWorkspace(sourceContent, workspaceUrl)
        logger.debug(">>> init done")
    return workspaceContent

# preprocesses the single string of contents and returns the following list: [words, tags]
#   where words is a list of words [w1, w2, w3,...]
#   tags is a dictionary that maps a tagger (e.g. stanford, nltk...) to a list of tags associated to the words list
#       e.g. tags = {
#                   "original" : [t1, t2, t3,...],
#                   ...
#                   }
def preprocessTerms(contents):
    splitContents = contents.split(" ")
    size = len(splitContents)
    words, tags = [], {"original": []}
    # ============== populate the words array and the tags dict =====================
    for i in range(size):  # iters thru term where term is <word>_<tag>
        term = splitContents[i]  # word_tag
        if "_" in term:
            splitTerm = term.split("_")
            word, originalTag = splitTerm[0], splitTerm[1]
            # add in a check the tag should be valid
            if validateTag(originalTag): # we only consider terms that have valid tags, as per this validator helper fn
               # separated words and tags into different arrays
                words.append(word)
                tags["original"].append(originalTag)
    tags["nltk"] = generateNLTKtags(words)
    # tags["spacy_sm"], tags["spacy_md"], tags["spacy_lg"] = generateSpacyTags(words)
    # ===============================================================================
    return [words, tags] # returns a list consisting of 1: a list and 2: a dictionary




# auto tags the confirmed singlish words and retuns the updated tags and updatedUncertainTagIndices. singlish words will have SFP tagging
# TODO: this needs to be changed because we assumed that the tag for every singlish word will be constant. This is
#       not true because things like "lah" can have multiple grammatical meanings.
def autoTagWords(scores, words, currentTags, currentUncertainTagIndices):
    updatedTags, updatedUncertainTagIndices = [], currentUncertainTagIndices
    print(f" SEE HERE scores length: {len(scores)}, words length: {len(words)},tags length: {len(currentTags)}")
    assert len(scores) == len(words) == len(currentTags)
    for idx in range(len(words)):
        word, score, updatedTag = words[idx], scores[idx], currentTags[idx]
        # autoTagsSinglish words:
        if (word in SinglishWords and score == -1):  # just being extra safe, could have just iterated thorugh scores without looking thru any of the words actually TODO: consider optimising this
            # autotag
            updatedTag = "SFP"
            # remove from uncertainTagIndices if the idx has been flagged
            if idx in currentUncertainTagIndices:
                updatedUncertainTagIndices.remove(idx)
        updatedTags.append(updatedTag)
    return words, updatedTags, updatedUncertainTagIndices


#TODO: replace the other output file shit.

# We create a new docx file because it's complicated to modify the input docx file because of the internal structure
# of docx. This isn't a problem because from a user perspective, the generated docx file looks exactly the same as the
# input docx file
def writeToDocx(scores, words, tagsDict):
    document = Document()
    p = document.add_paragraph("")

    for idx in range(len(words)):
        # init some variables:
        score = scores[idx]
        currentTag = tagsDict["original"][idx]
        word = words[idx]
        term = word + "_" +  currentTag

        # highlight terms based on their meanings:
        if score == SINGLISH_SCORE: # score of -1 had been assigned
            p.add_run(term).font.highlight_color = SINGLISH_HIGHLIGHT_COLOUR
        elif score == PSEUDOSINGLISH_SCORE: # score of 0 assigned
            p.add_run(term).font.highlight_color = PSEUDOSINGLISH_HIGHLIGHT_COLOUR
        elif score > PSEUDOSINGLISH_SCORE and score < THRESHOLD: # this didn't meet the min threshold
            p.add_run(term).font.highlight_color = UNCERTAIN_HIGHLIGHT_COLOUR
        else:
            p.add_run(term) # no issue in this
        p.add_run(" ")
    document.save(DESTINATION_FILE_PATH)
    return




# correctly assigns the POS to each word, returns a valid list of lists: [words, tags] where both words and tags are a list
def checkPOS(contents):
    words, tagsDict = preprocessTerms(contents)
    assert len(words) == len(tagsDict["original"]) == len(tagsDict["nltk"]), "words and tag arrays are not the same length in checkPOS()"
    originalTags, nltkTags =  tagsDict["original"], tagsDict["nltk"]
    scores = calculateScores(words, tagsDict)
    logger.info(f"\n Scores have been calculated. \n {scores} ")
    writeToDocx(scores, words, tagsDict)

    """ # code here uses multiple other libraries
    assert (len(words) == len(tagsDict["original"]) == len(tagsDict["nltk"]) == len(tagsDict["spacy_sm"])), "words and tag arrays are not the same length in checkPOS()"
    originalTags, nltkTags, spacyTags_sm, spacyTags_md, spacyTags_ls = \
        tagsDict["original"], tagsDict["nltk"], tagsDict["spacy_sm"], tagsDict["spacy_md"], tagsDict["spacy_lg"]
    scores = calculateScores(words, tagsDict)
    logger.info(f"\n Scores have been calculated. \n {scores} ")
    writeToDocx(scores, words, tagsDict) 
    """


    """ # code here involves manually asking for human input for the terms that have undesirable level of discrepancy 
    # amongst the various other nlp tokenizing models used (stanford, nltk...) 
        suspect_indices = detectDiscrepencies(scores, THRESHOLD)
        # indices ref to things with discrepancies:

        words, updatedTags, uncertainTagIndices = autoTagWords(scores, words, originalTags, suspect_indices)
        numberOfUncertainties = len(uncertainTagIndices)
        logger.info(f" THRESHOLD VALUE of {THRESHOLD} gives us {numberOfUncertainties} uncertainties that require human assistance out of {len(words)} words. \n {0 if len(words) == 0 else 100 * (numberOfUncertainties / len(words))} % of tagging done by Stanford Tagger needs to be reviewed by a human. ")
        finalisedTags = []
        for idx in range(
                len(words)
        ):  # asks for human to check only for the uncertain indices
            currentTag = updatedTags[idx]
            if idx in uncertainTagIndices:
                wordsBefore, wordsAfter = 0 if idx < 10 else (idx - 10), len(words) if idx >= len(words) - 10 else (
                            idx + 10)
                currentTerm = "" + words[idx] + "_" + currentTag
                # TODO: abstract this reference part out to reduce clutter here
                # generates some reference text to help determine the correct tag for that word:
                # ==================================================================================
                referenceText = "\n\t\t Here's the nearby text for reference:\n\n ================================ \n\t\t"
                for x in range(wordsBefore, wordsAfter):
                    referenceText += (("{" + words[x] + "}") if idx == x else words[x]) + " "
                # ==================================================================================
                finalisedTag = determineCorrectTag(currentTerm, referenceText)
                finalisedTags.append(finalisedTag)
            else:
                finalisedTags.append(currentTag)

        stats = f"\n\nThe human was involved for {str(numberOfUncertainties)} times for {len(words)} valid terms."
        return (words, finalisedTags, stats)
    """


# determines which scores are too low, and judges that as a discrepency
# returns a list of indices respective to the original array to reflect what needs to be prompted to human
# scores: the word list: [w1, w2, ...] will have associated scores for each word as a list: [s1, s2...]
# threshold: a minimum score for a single word. 
def detectDiscrepencies(scores, threshold):
    discrepencies = []
    for idx in range(len(scores)):
        score = scores[idx]
        if score < threshold:
            discrepencies.append(idx)
    return discrepencies  # list of index for humanPrompt to execute on

# asks human what the correct tag should be(<word>_<tag>) because a discrepancy has been found
# returns the correct tag for that particular word
def determineCorrectTag(term, referenceText):
    # referenceText = "\n\t\t Here's the nearby text for reference:\n\n ================================ \n\t\t"
    # for word in neighbouringWords:
    #     referenceText += word + " "
    print(referenceText)
    splitTerm = term.split("_")
    word, tag = splitTerm
    isValidTag = False
    while not isValidTag:
        prompt = f"Is this valid? \n \t [ {word} \n\t --------- \n\t       {tag} ] \n \t type Y or N "
        userInput = ""
        try:
            userInput = (input(prompt)).upper()
        except KeyboardInterrupt:
            # TODO: figure out how to do the pausing later
            logger.debug("DETECTED CTRL C")
        isValidTag = (userInput == "Y") and validateTag(tag)
        logging.info('Checked for the word {' + word + '} and the tag {' + tag + '}')
        if not isValidTag:
            showHelp()
            userInput = input("Enter valid tag:").upper()
            logging.info('Tag entered: ' + userInput + ' for the word ' + word)
            if userInput not in PosDictionary.keys():
                continue
            else:
                tag = PosDictionary.get(userInput)
        tagsChanged.append(tag)
    return tag


# compares tag arrs and gives a weighted sum of their matches based on their published frequencies.
def calculateWeightedSum(tagsDict, idx):
    stanfordTag = tagsDict["original"][idx]
    nltkTag = tagsDict["nltk"][idx]
    #spacyTag_sm = tagsDict["spacy_sm"][idx]
    #spacyTag_md = tagsDict["spacy_md"][idx]
    #spacyTag_lg = tagsDict["spacy_lg"][idx]

    c0 = 1 * (PublishedAccuracies["stanford"])  # because the original tags were generated from stanford, it's guaranteed to be a match, this is equivalent to (1 if stanfordTag == stanfordTag else 0) * (PublishedAccuracies["stanford"])
    c1 = (1 if stanfordTag == nltkTag else 0) * (PublishedAccuracies["nltk"])
    #c2 = (1 if stanfordTag == spacyTag_sm else 0) * (PublishedAccuracies["spacy_sm"])
    #c3 = (1 if stanfordTag == spacyTag_md else 0) * (PublishedAccuracies["spacy_md"])
    #c4 = (1 if stanfordTag == spacyTag_lg else 0) * (PublishedAccuracies["spacy_lg"])
    #return c0 + c1 + c2 + c3 + c4
    return c0 + c1


# input: dictionary of generated Tags
#  returns a list of scores
# if it's a confirmed to be a singlish wrod, then we set the score as -1
# if it's a pseudoSinglish word, e.g. "one" then we set the score to be 0

# TODO: make this generalised once we use more than 1 reference models
def calculateScores(words, generatedTags):
    size = len(generatedTags["original"])
    assert size == len(generatedTags["nltk"])
    scores = []
    for idx in range(size):
        word = words[idx].upper()
        if word in SinglishWords:
            score = SINGLISH_SCORE  # looking at scores array, we can say that - 1 means safe to autocheck without asking the human
        elif word in PseudoSinglishWords:
            score = PSEUDOSINGLISH_SCORE  # means there's definitely gonna be some discrepancy
        else: # actual weighted score calculation
            score = calculateWeightedSum(generatedTags, idx)
        scores.append(score)
    return scores


# cleans up the program upon a signal interruption by
# 1: appending checkedWords and checkedTags to the output file
# 2: overwriting the workspace file with the uncheckedWords and uncheckedTags
def cleanup(checkedWords, checkedTags, uncheckedWords, uncheckedTags):
    writeToOutputFile(checkedWords, checkedTags)
    remainingContent = ""
    for i in range(len(uncheckedWords)):
        remainingContent += str(uncheckedWords[i]) + "_" + str(uncheckedTags[i]) + " "
    writeToWorkspace(
        remainingContent, workspaceUrl="workspace_PoSTrialText.txt"
    )  # TODO: make everything into a class so that can set the filenames as a class level variable after init
    os.kill(os.getpid(), signal.SIGINT)


# writes content to a workspace file. Intentionally overwrites if there's an existing file:
def writeToWorkspace(content, workspaceUrl):
    workspaceFile = open(workspaceUrl, "w+")
    # copy over to the new file name:
    workspaceFile.write(content)
    workspaceFile.close()
    logger.debug(">>> wrote to Workspace")
    return content


# Returns a list of filenames based on what input files are placed inside the /Inputs dir
def getInputFilepaths():
    inputLocation = os.path.join(os.getcwd(), "Inputs")
    listOfFiles = os.listdir(inputLocation)
    filenames =[]
    for entry in listOfFiles:
        fullPath = os.path.join(inputLocation, entry)
        filenames.append(fullPath)
    return filenames

# appends to non-existing / pre-existing output file
def writeToOutputFile(fileTitle, words, tags):
    outputFileUrl, submissionFileUrl =\
        os.path.join(os.getcwd(),"Outputs",fileTitle + ".txt") , os.path.join(os.getcwd(),"Submissions",fileTitle + ".txt")

    outputFile, submissionFile = open(outputFileUrl, "a+"), open(submissionFileUrl, "a+")
    size = len(words)
    date = datetime.now().isoformat()
    outputString, submissionString = f"\n===================== { datetime.now().isoformat()} STARTING LINE =============================\n",f"\n========= {datetime.now().isoformat()} START SUBMISSION==========\n"
    for i in range(size):
        word, tag = words[i], tags[i]
        outputEntry, submissionEntry = "[" + str(word) + "_" + str(tag) + "]", str(word) + "_" + str(tag) + " "
        outputString += outputEntry + "\n"
        submissionString += submissionEntry
    outputFile.write(outputString + f"\n======= {datetime.now().isoformat()} ============\n")
    submissionFile.write(submissionString + f"\n======= {datetime.now().isoformat()} ============\n")
    outputFile.close(); submissionFile.close()

def validateTag(tag):
    return tag in PosDictionary.values()


# tagging accuracies of models when input is standard English
# assumptions:
# -  we have ignored that fact that the testing done to get these accuracy values may/may not have been from the same
#   benchmarking source.
# -
PublishedAccuracies = {
    "stanford": 0.9697,   # https://tinyurl.com/stanfordTagger-accuracy
    "nltk": 0.94, # TODO: add link to the published accuracy value for ntlk
    "spacy_sm": 0.966, # https://github.com/explosion/spacy-models/releases/en_core_web_sm-1.2.0
    "spacy_md": 0.967, #https://github.com/explosion/spacy-models/releases/en_core_web_md-1.2.0
    "spacy_lg": 0.9722, #https://github.com/explosion/spacy-models/releases/tag/en_core_web_lg-2.3.1
}



# TODO:  init this dictionary properly please.

PosDictionary = {}
PosDictionary["1"] = "CC"
PosDictionary["2"] = "CD"
PosDictionary["3"] = "DT"
PosDictionary["4"] = "EX"
PosDictionary["5"] = "FW"
PosDictionary["6"] = "IN"
PosDictionary["7"] = "JJ"
PosDictionary["8"] = "JJR"
PosDictionary["9"] = "JJS"
PosDictionary["10"] = "LS"
PosDictionary["11"] = "MD"
PosDictionary["12"] = "NN"
PosDictionary["13"] = "NNS"
PosDictionary["14"] = "NNP"
PosDictionary["15"] = "NNPS"
PosDictionary["16"] = "PDT"
PosDictionary["17"] = "POS"
PosDictionary["18"] = "PRP"
PosDictionary["19"] = "PRP$"
PosDictionary["20"] = "RB"
PosDictionary["21"] = "RBR"
PosDictionary["22"] = "RBS"
PosDictionary["23"] = "RP"
PosDictionary["24"] = "SYM"
PosDictionary["25"] = "TO"
PosDictionary["26"] = "UH"
PosDictionary["27"] = "VB"
PosDictionary["28"] = "VBD"
PosDictionary["29"] = "VBG"
PosDictionary["30"] = "VBN"
PosDictionary["31"] = "VBP"
PosDictionary["32"] = "VBZ"
PosDictionary["33"] = "WDT"
PosDictionary["34"] = "WP"
PosDictionary["35"] = "WP$"
PosDictionary["36"] = "WRB"
PosDictionary["37"] = "SFP"


def showHelp():
    line = "___________________________________________________\n"
    keys, values = list(PosDictionary.keys()), list(PosDictionary.values())
    message = "\n----- {enter number representing the tag} -------- \n" + line
    leftPtr, rightPtr = 0, len(keys) - 1
    while (rightPtr >= leftPtr):
        leftKey, leftValue = keys[leftPtr], values[leftPtr]
        rightKey, rightValue = keys[rightPtr], values[rightPtr]
        if leftPtr != rightPtr:
            if rightPtr >= 25:
                message += (
                    f"{leftKey}: {leftValue} \t\t\t {rightKey}: {rightValue} \n")  # this is just to prettify the printing, purely aesthetic
            else:
                message += (f"{leftKey}: {leftValue} \t\t {rightKey}: {rightValue} \n")
        else:
            message += (f"{rightKey}: {rightValue} \n")
        rightPtr -= 1;
        leftPtr += 1
    message += line
    print(message)


# def outputLoggingFile(changedTags):


endingGreeting = f"""
================================================================================================ \n
{f.renderText("I am so done")}
 \n
================================================================================================

"""

if __name__ == "__main__":
    main()
